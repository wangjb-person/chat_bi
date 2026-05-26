"""
知识库文件解析：按后缀选择解析器，输出纯文本并推断切块策略。
"""

from __future__ import annotations

import csv
import io
from collections.abc import Callable
from dataclasses import dataclass
from pathlib import Path
from typing import List

import pandas as pd

from chatbi.services.kb_chunking import ChunkProfile, FileKind, infer_chunk_profile

SUPPORTED_SUFFIXES = frozenset(
    {
        ".txt",
        ".md",
        ".markdown",
        ".csv",
        ".pdf",
        ".xlsx",
        ".xls",
        ".docx",
    }
)

_SUFFIX_TO_KIND: dict[str, FileKind] = {
    ".txt": FileKind.TXT,
    ".md": FileKind.MARKDOWN,
    ".markdown": FileKind.MARKDOWN,
    ".csv": FileKind.CSV,
    ".pdf": FileKind.PDF,
    ".xlsx": FileKind.EXCEL,
    ".xls": FileKind.EXCEL,
    ".docx": FileKind.WORD,
}


@dataclass(frozen=True)
class ParsedKbDocument:
    """解析结果：正文 + 类型 + 推荐切块策略。"""

    text: str
    file_kind: FileKind
    chunk_profile: ChunkProfile
    suffix: str


def normalize_suffix(filename: str) -> str:
    """提取并规范化文件后缀（小写，含点）。"""
    return Path(filename or "").suffix.lower()


def is_supported_upload(filename: str) -> bool:
    """上传文件名后缀是否在支持列表内。"""
    return normalize_suffix(filename) in SUPPORTED_SUFFIXES


def supported_suffixes_hint() -> str:
    """支持格式说明，用于错误提示。"""
    return ", ".join(sorted(SUPPORTED_SUFFIXES))


def parse_upload(raw_bytes: bytes, filename: str) -> ParsedKbDocument:
    """解析上传二进制为纯文本，并推断切块策略。"""
    if not raw_bytes:
        raise ValueError("文件内容为空")

    suffix = normalize_suffix(filename)
    if suffix not in SUPPORTED_SUFFIXES:
        raise ValueError(
            f"不支持该文件类型 {suffix or '(无后缀)'}，支持：{supported_suffixes_hint()}"
        )

    kind = _SUFFIX_TO_KIND[suffix]
    parser = _BYTE_PARSERS.get(kind, _decode_text)
    text = parser(raw_bytes, suffix=suffix).strip()
    if not text:
        raise ValueError("解析后无有效文本，请检查文件是否为空或是否为扫描版 PDF")

    return _build_parsed(text, kind=kind, suffix=suffix)


def parse_paste(content: str, *, filename: str = "paste.txt") -> ParsedKbDocument:
    """解析粘贴文本（按文件名后缀辅助推断类型）。"""
    text = (content or "").strip()
    if not text:
        raise ValueError("文档内容不能为空")

    suffix = normalize_suffix(filename) or ".txt"
    kind = _SUFFIX_TO_KIND.get(suffix, FileKind.TXT)
    return _build_parsed(text, kind=kind, suffix=suffix)


def _build_parsed(text: str, *, kind: FileKind, suffix: str) -> ParsedKbDocument:
    """组装解析结果并推断切块策略。"""
    profile = infer_chunk_profile(text, file_kind=kind, suffix=suffix)
    return ParsedKbDocument(
        text=text,
        file_kind=kind,
        chunk_profile=profile,
        suffix=suffix,
    )


def _decode_text(raw: bytes, *, suffix: str = "") -> str:
    """解码纯文本（UTF-8 / GBK 等，可选 chardet）。"""
    del suffix
    for enc in ("utf-8", "utf-8-sig", "gbk", "gb18030"):
        try:
            return raw.decode(enc)
        except UnicodeDecodeError:
            continue
    try:
        import chardet

        enc = chardet.detect(raw).get("encoding") or "utf-8"
        return raw.decode(enc, errors="replace")
    except ImportError:
        return raw.decode("utf-8", errors="replace")


def _parse_csv(raw: bytes, *, suffix: str = "") -> str:
    """CSV → 制表符分隔的纯文本行。"""
    del suffix
    text = _decode_text(raw)
    sample = text[:4096]
    try:
        dialect = csv.Sniffer().sniff(sample, delimiters=",;\t|")
    except csv.Error:
        dialect = csv.excel
    rows = list(csv.reader(io.StringIO(text), dialect))
    if not rows:
        return text
    return "\n".join("\t".join(cell.strip() for cell in row) for row in rows)


def _parse_pdf(raw: bytes, *, suffix: str = "") -> str:
    """PDF 逐页抽取文本。"""
    del suffix
    try:
        from pypdf import PdfReader
    except ImportError as e:
        raise ValueError("PDF 解析依赖未安装，请执行: pip install pypdf") from e

    reader = PdfReader(io.BytesIO(raw))
    parts = [
        (page.extract_text() or "").strip()
        for page in reader.pages
        if (page.extract_text() or "").strip()
    ]
    if not parts:
        raise ValueError(
            "PDF 未提取到文本，可能为扫描件，请先 OCR 或转为 Word/txt 后上传"
        )
    return "\n\n".join(parts)


def _parse_excel(raw: bytes, *, suffix: str = "") -> str:
    """Excel 各工作表转为带标题的文本块。"""
    engine = "xlrd" if suffix == ".xls" else None
    try:
        xl = pd.ExcelFile(io.BytesIO(raw), engine=engine)
    except ImportError as e:
        hint = "openpyxl" if suffix == ".xlsx" else "xlrd"
        raise ValueError(f"Excel 解析依赖未安装，请执行: pip install {hint}") from e
    except Exception as e:
        raise ValueError(f"Excel 解析失败：{e}") from e

    sections: List[str] = []
    for sheet in xl.sheet_names:
        df = pd.read_excel(xl, sheet_name=sheet, dtype=str).fillna("")
        if df.empty:
            continue
        table = df.to_csv(index=False, sep="\t").strip()
        sections.append(f"## 工作表: {sheet}\n{table}")

    if not sections:
        raise ValueError("Excel 中无有效数据")
    return "\n\n".join(sections)


def _parse_docx(raw: bytes, *, suffix: str = "") -> str:
    """Word docx：段落与表格单元格。"""
    del suffix
    try:
        from docx import Document
    except ImportError as e:
        raise ValueError("Word 解析依赖未安装，请执行: pip install python-docx") from e

    doc = Document(io.BytesIO(raw))
    parts: List[str] = [
        (para.text or "").strip()
        for para in doc.paragraphs
        if (para.text or "").strip()
    ]
    for table in doc.tables:
        for row in table.rows:
            cells = [(c.text or "").strip() for c in row.cells]
            if any(cells):
                parts.append("\t".join(cells))
    if not parts:
        raise ValueError("Word 文档中无有效文本")
    return "\n\n".join(parts)


_BYTE_PARSERS: dict[FileKind, Callable[[bytes], str]] = {
    FileKind.TXT: _decode_text,
    FileKind.MARKDOWN: _decode_text,
    FileKind.CSV: _parse_csv,
    FileKind.PDF: _parse_pdf,
    FileKind.EXCEL: _parse_excel,
    FileKind.WORD: _parse_docx,
}
